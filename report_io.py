"""Формирование Word-отчёта с результатами расчёта."""

import math
from io import BytesIO

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def _custom_round(x, min_decimals=None):
    """Как в app.custom_round: 3 значащие цифры; |x|>=100 — до целого."""
    x = float(x)
    abs_x = abs(x)

    if abs_x >= 100:
        decimals = 0
    elif abs_x == 0:
        decimals = 0
    else:
        text = np.format_float_positional(x, precision=3, fractional=False, trim="-")
        decimals = len(text.split(".")[1]) if "." in text else 0

    if min_decimals is not None:
        decimals = max(decimals, int(min_decimals))

    return f"{x:.{decimals}f}"


def _format_cell(val, min_decimals=None):
    if val is None or (isinstance(val, str) and val.strip() == ""):
        return ""
    if isinstance(val, str):
        return val
    try:
        num = float(val)
    except (TypeError, ValueError):
        return str(val)
    if not math.isfinite(num):
        return "Не существует"
    return _custom_round(num, min_decimals=min_decimals)


def _format_parameters_df(df, data_precision=None, lang="ru"):
    """Округление метрик как на сайте (style_dataframe_html)."""
    from i18n import t as _t

    out = df.copy()
    out.index.name = _t("distribution", lang=lang)
    out.columns.name = None
    # Mean (0), MAE (4), maxE (5) — в единицах исходного ряда
    data_scale_cols = {0, 4, 5}
    for j, col in enumerate(out.columns):
        min_decimals = data_precision if j in data_scale_cols else None
        out[col] = out[col].map(
            lambda v, md=min_decimals: _format_cell(v, min_decimals=md)
        )
    return out


def _add_dataframe_table(
    doc,
    df,
    include_index=True,
    index_label=None,
    show_header=True,
    default_index_label="Distribution",
):
    """Добавляет таблицу DataFrame в документ."""
    df = df.copy()
    if include_index:
        label = index_label
        if label is None:
            label = df.index.name if df.index.name not in (None, "") else None
        if label is None:
            label = default_index_label
        df.index.name = label
        df = df.reset_index(drop=False)
        first = df.columns[0]
        if str(first) == "index":
            df = df.rename(columns={first: label})
    else:
        df = df.reset_index(drop=True)

    n_cols = len(df.columns)
    if show_header:
        table = doc.add_table(rows=1, cols=n_cols)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col)
    else:
        table = doc.add_table(rows=0, cols=n_cols)
        table.style = "Table Grid"

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            cells[i].text = "" if pd.isna(val) else str(val)


def _make_series_chart_png(
    series_df, values_col, index_col, title, no_group_marker
):
    """PNG графика хода значений (matplotlib)."""
    fig, ax = plt.subplots(figsize=(12, 6))
    x = series_df[index_col] if index_col != no_group_marker else series_df.index
    y = series_df[values_col]
    ax.plot(x, y, linewidth=1.5)
    ax.set_ylabel(values_col, fontsize=12)
    ax.set_title(title, fontsize=14)
    ax.tick_params(axis="x", labelsize=11)
    ax.tick_params(axis="y", labelsize=11)
    if index_col != no_group_marker:
        ax.set_xlabel(index_col, fontsize=12)
    fig.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def build_results_docx(
    series_df,
    data,
    values_col,
    index_col,
    curve_png=None,
    parameters_df=None,
    quantiles_df=None,
    data_precision=None,
    lang="ru",
    no_group_marker="__no_group__",
):
    """
    Word: эмпирический ряд, график хода, кривые, метрики, таблица обеспеченностей.
    Возвращает bytes готового .docx.
    """
    from i18n import t as _t

    doc = Document()
    title = doc.add_heading(_t("docx_title", lang=lang), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(_t("docx_empirical", lang=lang), level=1)
    ensurance_col = "Обеспеченность P, %"
    if index_col != no_group_marker:
        table_df = data[
            [
                index_col,
                values_col,
                ensurance_col,
                values_col + " (P)",
                index_col + " (P)",
            ]
        ].copy()
    else:
        table_df = data[[values_col, ensurance_col, values_col + " (P)"]].copy()
    table_df = table_df.rename(
        columns={ensurance_col: _t("ensurance_p_pct", lang=lang)}
    )
    _add_dataframe_table(
        doc,
        table_df,
        include_index=True,
        index_label=_t("rank", lang=lang),
    )

    doc.add_heading(_t("docx_series_chart", lang=lang), level=1)
    chart_png = _make_series_chart_png(
        series_df,
        values_col,
        index_col,
        title=_t("chart_series_title", lang=lang),
        no_group_marker=no_group_marker,
    )
    doc.add_picture(chart_png, width=Cm(16))

    if curve_png:
        doc.add_heading(_t("docx_curves", lang=lang), level=1)
        doc.add_picture(BytesIO(curve_png), width=Cm(16))

    if parameters_df is not None and not parameters_df.empty:
        doc.add_heading(_t("docx_metrics", lang=lang), level=1)
        params_fmt = _format_parameters_df(
            parameters_df, data_precision=data_precision, lang=lang
        )
        _add_dataframe_table(
            doc,
            params_fmt,
            include_index=True,
            index_label=_t("distribution", lang=lang),
            show_header=True,
        )

    if quantiles_df is not None and not quantiles_df.empty:
        doc.add_heading(_t("docx_quantiles", lang=lang), level=1)
        q = quantiles_df.copy()
        q.columns.name = None
        q.index.name = None
        _add_dataframe_table(
            doc,
            q,
            include_index=True,
            index_label="",
            show_header=False,
        )

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None:
                run.font.size = Pt(11)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
